"""
generar_documento.py — Genera el Word de entrega (Guía 4) con imágenes.
Coloca las capturas en la carpeta Capturas/ con los nombres indicados y ejecuta:
    python generar_documento.py

Nombres de archivos esperados en Capturas/ (.png o .jpg):
    01_estructura.png       → Carpeta del proyecto en VS Code
    02_menu.png             → Menú principal (python main.py)
    03_crear_ok.png         → Crear producto exitoso
    04_crear_error.png      → Crear producto - ID duplicado
    05_leer_ok.png          → Leer producto por ID exitoso
    07_listar.png           → Listar todos los productos
    08_actualizar_ok.png    → Actualizar producto exitoso
    09_actualizar_error.png → Actualizar producto - ID no existe
    10_eliminar_ok.png      → Eliminar producto exitoso
    11_eliminar_error.png   → Eliminar producto - ID no existe
    12_tests.png            → 22 pruebas unitarias OK
    13_sonarqube.png        → Resultados SonarCloud/SonarQube
    14_github.png           → Repositorio GitHub
"""
import os
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CAPTURAS_DIR = "Capturas"

# ── helpers ──────────────────────────────────────────────────────────────── #

def find_image(name_without_ext: str) -> str | None:
    """Busca la imagen en Capturas/ aceptando .png, .jpg o .jpeg."""
    for ext in (".png", ".jpg", ".jpeg"):
        path = os.path.join(CAPTURAS_DIR, name_without_ext + ext)
        if os.path.exists(path):
            return path
    return None


def insert_image(doc: Document, filename_base: str, label: str, width: float = 6.0):
    """Inserta la imagen si existe; si no, pone un marcador de texto."""
    path = find_image(filename_base)
    if path:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = cap.add_run(f"Figura: {label}")
        run2.font.italic = True
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[ CAPTURA PENDIENTE: {label} | guardar como {filename_base}.png ]")
        run.font.color.rgb = RGBColor(0xBB, 0x00, 0x00)
        run.font.italic = True
        run.font.size = Pt(10)
    doc.add_paragraph()


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_code_block(doc: Document, code: str):
    style_name = "CodeBlock"
    if style_name not in [s.name for s in doc.styles]:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Courier New"
        style.font.size = Pt(8)
        pf = style.paragraph_format
        pf.left_indent = Cm(0.5)
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
    for line in code.splitlines():
        p = doc.add_paragraph(line if line else " ", style=style_name)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)


def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def body(doc, text):
    doc.add_paragraph(text)


# ── código fuente ─────────────────────────────────────────────────────────── #

PRODUCTOS_PY = '''\
import json, os

class ProductoCRUD:
    def __init__(self, archivo: str = None):
        self._productos: dict = {}
        self._archivo = archivo
        if archivo and os.path.exists(archivo):
            self._cargar()

    # CREATE
    def crear(self, id, nombre, descripcion, precio, cantidad):
        if id in self._productos:
            raise ValueError(f"Ya existe un producto con id={id}")
        if precio < 0: raise ValueError("El precio no puede ser negativo")
        if cantidad < 0: raise ValueError("La cantidad no puede ser negativa")
        producto = {"id": id, "nombre": nombre, "descripcion": descripcion,
                    "precio": precio, "cantidad": cantidad}
        self._productos[id] = producto
        self._guardar()
        return dict(producto)

    # READ
    def leer(self, id):
        if id not in self._productos:
            raise KeyError(f"No existe producto con id={id}")
        return dict(self._productos[id])

    def leer_todos(self):
        return [dict(p) for p in self._productos.values()]

    # UPDATE
    def actualizar(self, id, **campos):
        if id not in self._productos:
            raise KeyError(f"No existe producto con id={id}")
        campos_validos = {"nombre", "descripcion", "precio", "cantidad"}
        invalidos = set(campos) - campos_validos
        if invalidos: raise ValueError(f"Campos no válidos: {invalidos}")
        if "precio" in campos and campos["precio"] < 0:
            raise ValueError("El precio no puede ser negativo")
        if "cantidad" in campos and campos["cantidad"] < 0:
            raise ValueError("La cantidad no puede ser negativa")
        self._productos[id].update(campos)
        self._guardar()
        return dict(self._productos[id])

    # DELETE
    def eliminar(self, id):
        if id not in self._productos:
            raise KeyError(f"No existe producto con id={id}")
        eliminado = self._productos.pop(id)
        self._guardar()
        return dict(eliminado)

    def _guardar(self):
        if self._archivo:
            with open(self._archivo, "w", encoding="utf-8") as f:
                json.dump(list(self._productos.values()), f, ensure_ascii=False, indent=2)

    def _cargar(self):
        with open(self._archivo, "r", encoding="utf-8") as f:
            datos = json.load(f)
        self._productos = {p["id"]: p for p in datos}
'''

TESTS_PY = '''\
import unittest
from productos import ProductoCRUD

class TestCrear(unittest.TestCase):
    def setUp(self): self.crud = ProductoCRUD()

    def test_crear_producto_retorna_datos_correctos(self):
        p = self.crud.crear(1, "Laptop", "Computadora portátil", 2500.0, 10)
        self.assertEqual(p["id"], 1); self.assertEqual(p["precio"], 2500.0)

    def test_crear_producto_queda_almacenado(self):
        self.crud.crear(2, "Mouse", "Periférico", 50.0, 25)
        self.assertEqual(self.crud.leer(2)["nombre"], "Mouse")

    def test_crear_id_duplicado_lanza_ValueError(self):
        self.crud.crear(1, "Laptop", "Desc", 2500.0, 10)
        with self.assertRaises(ValueError): self.crud.crear(1, "Mouse", "Desc2", 50.0, 5)

    def test_crear_precio_negativo_lanza_ValueError(self):
        with self.assertRaises(ValueError): self.crud.crear(3, "Teclado", "Desc", -100.0, 5)

    def test_crear_cantidad_negativa_lanza_ValueError(self):
        with self.assertRaises(ValueError): self.crud.crear(4, "Monitor", "Desc", 500.0, -1)

class TestLeer(unittest.TestCase):
    def setUp(self):
        self.crud = ProductoCRUD()
        self.crud.crear(1, "Laptop", "Computadora portátil", 2500.0, 10)
        self.crud.crear(2, "Mouse", "Periférico", 50.0, 25)

    def test_leer_producto_existente(self):
        self.assertEqual(self.crud.leer(1)["nombre"], "Laptop")

    def test_leer_todos_devuelve_lista_completa(self):
        self.assertEqual(len(self.crud.leer_todos()), 2)

    def test_leer_todos_vacio_devuelve_lista_vacia(self):
        self.assertEqual(ProductoCRUD().leer_todos(), [])

    def test_leer_id_inexistente_lanza_KeyError(self):
        with self.assertRaises(KeyError): self.crud.leer(999)

    def test_leer_devuelve_copia_independiente(self):
        p = self.crud.leer(1); p["nombre"] = "Cambiado"
        self.assertEqual(self.crud.leer(1)["nombre"], "Laptop")

class TestActualizar(unittest.TestCase):
    def setUp(self):
        self.crud = ProductoCRUD()
        self.crud.crear(1, "Laptop", "Computadora portátil", 2500.0, 10)

    def test_actualizar_precio_y_cantidad(self):
        p = self.crud.actualizar(1, precio=2000.0, cantidad=8)
        self.assertEqual(p["precio"], 2000.0); self.assertEqual(p["cantidad"], 8)

    def test_actualizar_nombre(self):
        self.assertEqual(self.crud.actualizar(1, nombre="Laptop Pro")["nombre"], "Laptop Pro")

    def test_actualizar_descripcion(self):
        self.assertEqual(self.crud.actualizar(1, descripcion="Nueva")["descripcion"], "Nueva")

    def test_actualizar_producto_inexistente_lanza_KeyError(self):
        with self.assertRaises(KeyError): self.crud.actualizar(999, nombre="X")

    def test_actualizar_precio_negativo_lanza_ValueError(self):
        with self.assertRaises(ValueError): self.crud.actualizar(1, precio=-50.0)

    def test_actualizar_cantidad_negativa_lanza_ValueError(self):
        with self.assertRaises(ValueError): self.crud.actualizar(1, cantidad=-5)

    def test_actualizar_campo_invalido_lanza_ValueError(self):
        with self.assertRaises(ValueError): self.crud.actualizar(1, color="rojo")

class TestEliminar(unittest.TestCase):
    def setUp(self):
        self.crud = ProductoCRUD()
        self.crud.crear(1, "Laptop", "Computadora portátil", 2500.0, 10)
        self.crud.crear(2, "Mouse", "Periférico", 50.0, 25)

    def test_eliminar_producto_existente(self):
        self.assertEqual(self.crud.eliminar(1)["id"], 1)
        self.assertEqual(len(self.crud.leer_todos()), 1)

    def test_eliminar_producto_ya_no_accesible(self):
        self.crud.eliminar(1)
        with self.assertRaises(KeyError): self.crud.leer(1)

    def test_eliminar_reduce_conteo(self):
        self.crud.eliminar(1); self.assertEqual(len(self.crud.leer_todos()), 1)

    def test_eliminar_producto_inexistente_lanza_KeyError(self):
        with self.assertRaises(KeyError): self.crud.eliminar(999)

    def test_eliminar_mismo_producto_dos_veces_lanza_KeyError(self):
        self.crud.eliminar(2)
        with self.assertRaises(KeyError): self.crud.eliminar(2)

if __name__ == "__main__":
    unittest.main(verbosity=2)
'''

# ── construcción del documento ───────────────────────────────────────────── #

def build_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # ── PORTADA ──────────────────────────────────────────────────────────── #
    for _ in range(3): doc.add_paragraph()

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("INSTITUTO TECNOLÓGICO METROPOLITANO — ITM")
    r.bold = True; r.font.size = Pt(14)

    doc.add_paragraph()
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run("Aplicaciones y Servicios Web  |  Código: 580202009").font.size = Pt(12)

    for _ in range(2): doc.add_paragraph()
    t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("GUÍA DE TRABAJO No. 4 — TALLER EVALUATIVO 20%")
    r2.bold = True; r2.font.size = Pt(16)

    doc.add_paragraph()
    s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = s2.add_run("Implementación de un CRUD en Python con Pruebas Unitarias")
    r3.italic = True; r3.font.size = Pt(13)

    for _ in range(2): doc.add_paragraph()

    tbl = doc.add_table(rows=5, cols=2); tbl.style = "Table Grid"
    labels = ["Estudiante(s):", "Programa:", "Semestre:", "Fecha de entrega:", "Docente:"]
    values = ["[NOMBRE COMPLETO]", "[PROGRAMA ACADÉMICO]", "[SEMESTRE]",
              datetime.date.today().strftime("%d de %B de %Y"),
              "Jonathan Sánchez Giraldo"]
    for i, (lbl, val) in enumerate(zip(labels, values)):
        set_cell_bg(tbl.rows[i].cells[0], "D9E1F2")
        tbl.rows[i].cells[0].paragraphs[0].add_run(lbl).bold = True
        tbl.rows[i].cells[1].paragraphs[0].add_run(val)

    doc.add_page_break()

    # ── 1. INTRODUCCIÓN ──────────────────────────────────────────────────── #
    heading(doc, "1. Introducción")
    body(doc,
         "El presente documento describe el desarrollo de un sistema CRUD (Create, Read, "
         "Update, Delete) para la gestión de un registro de productos, implementado en "
         "Python. El proyecto cumple con los requisitos de la Guía de Trabajo No. 4 de "
         "la asignatura Aplicaciones y Servicios Web: lenguaje Python, almacenamiento en "
         "memoria con persistencia en JSON, pruebas unitarias por cada operación y "
         "publicación del código en GitHub con análisis de calidad en SonarCloud.")

    # ── 2. ESTRUCTURA DEL PROYECTO ───────────────────────────────────────── #
    heading(doc, "2. Configuración Inicial y Estructura del Proyecto")
    body(doc, "El proyecto contiene los siguientes archivos:")

    tbl2 = doc.add_table(rows=4, cols=2); tbl2.style = "Table Grid"
    set_cell_bg(tbl2.rows[0].cells[0], "D9E1F2")
    set_cell_bg(tbl2.rows[0].cells[1], "D9E1F2")
    tbl2.rows[0].cells[0].paragraphs[0].add_run("Archivo").bold = True
    tbl2.rows[0].cells[1].paragraphs[0].add_run("Descripción").bold = True
    filas = [
        ("productos.py",      "Clase ProductoCRUD: lógica de negocio CRUD + persistencia JSON."),
        ("main.py",           "Menú interactivo en consola. Guarda datos en productos.json."),
        ("test_productos.py", "22 pruebas unitarias (unittest) para cada operación CRUD."),
    ]
    for i, (a, d) in enumerate(filas, 1):
        tbl2.rows[i].cells[0].paragraphs[0].add_run(a).font.name = "Courier New"
        tbl2.rows[i].cells[1].paragraphs[0].add_run(d)

    doc.add_paragraph()
    heading(doc, "2.1 Captura — Estructura del proyecto en VS Code", level=2)
    insert_image(doc, "01_estructura", "Estructura del proyecto en VS Code Explorer")

    # ── 3. MODELO DE DATOS ───────────────────────────────────────────────── #
    heading(doc, "3. Modelo de Datos — Producto")
    tbl3 = doc.add_table(rows=6, cols=3); tbl3.style = "Table Grid"
    for j, h in enumerate(["Campo", "Tipo", "Descripción"]):
        set_cell_bg(tbl3.rows[0].cells[j], "D9E1F2")
        tbl3.rows[0].cells[j].paragraphs[0].add_run(h).bold = True
    campos_data = [
        ("id",          "int",   "Identificador único del producto."),
        ("nombre",      "str",   "Nombre comercial del producto."),
        ("descripcion", "str",   "Breve descripción del producto."),
        ("precio",      "float", "Precio unitario (≥ 0)."),
        ("cantidad",    "int",   "Unidades disponibles en inventario (≥ 0)."),
    ]
    for i, (c, t, d) in enumerate(campos_data, 1):
        tbl3.rows[i].cells[0].paragraphs[0].add_run(c).font.name = "Courier New"
        tbl3.rows[i].cells[1].paragraphs[0].add_run(t).font.name = "Courier New"
        tbl3.rows[i].cells[2].paragraphs[0].add_run(d)

    # ── 4. IMPLEMENTACIÓN CRUD ───────────────────────────────────────────── #
    heading(doc, "4. Implementación del CRUD")
    heading(doc, "4.1 Código fuente — productos.py", level=2)
    body(doc, "Clase ProductoCRUD con todas las operaciones y persistencia en JSON:")
    add_code_block(doc, PRODUCTOS_PY)

    heading(doc, "4.2 Descripción de operaciones", level=2)
    ops = [
        ("crear",       "Valida id único, precio ≥ 0 y cantidad ≥ 0. Lanza ValueError si falla."),
        ("leer",        "Retorna copia del producto por id. Lanza KeyError si no existe."),
        ("leer_todos",  "Retorna lista completa de productos. Lista vacía si no hay ninguno."),
        ("actualizar",  "Modifica campos permitidos. Valida existencia y valores numéricos."),
        ("eliminar",    "Extrae el producto del diccionario y lo devuelve. Lanza KeyError si no existe."),
    ]
    for nombre_op, desc in ops:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{nombre_op}: ").bold = True
        p.add_run(desc)

    # ── 5. CAPTURAS CRUD ─────────────────────────────────────────────────── #
    doc.add_page_break()
    heading(doc, "5. Evidencia de Ejecución — Operaciones CRUD")

    heading(doc, "5.1 Menú principal", level=2)
    body(doc, "Ejecución de python main.py mostrando el menú interactivo.")
    insert_image(doc, "02_menu", "Menú principal — python main.py")

    heading(doc, "5.2 Crear producto — caso exitoso", level=2)
    body(doc, "Se ingresa un producto nuevo con todos sus campos. El sistema responde con [OK].")
    insert_image(doc, "03_crear_ok", "Crear producto exitoso")

    heading(doc, "5.3 Crear producto — caso de error (ID duplicado)", level=2)
    body(doc, "Se intenta crear un producto con un ID ya registrado. El sistema responde con [ERROR].")
    insert_image(doc, "04_crear_error", "Crear producto — ID duplicado")

    heading(doc, "5.4 Leer producto por ID — caso exitoso", level=2)
    body(doc, "Se consulta un producto existente por su ID y se muestran sus datos.")
    insert_image(doc, "05_leer_ok", "Leer producto por ID exitoso")

    heading(doc, "5.5 Listar todos los productos", level=2)
    body(doc, "Se listan todos los productos actualmente registrados en el sistema.")
    insert_image(doc, "07_listar", "Listar todos los productos")

    heading(doc, "5.6 Actualizar producto — caso exitoso", level=2)
    body(doc, "Se modifican uno o más campos de un producto existente. El sistema confirma con [OK].")
    insert_image(doc, "08_actualizar_ok", "Actualizar producto exitoso")

    heading(doc, "5.7 Actualizar producto — caso de error (ID no existe)", level=2)
    body(doc, "Se intenta actualizar un producto con ID inexistente. El sistema responde con [ERROR].")
    insert_image(doc, "09_actualizar_error", "Actualizar producto — ID inexistente")

    heading(doc, "5.8 Eliminar producto — caso exitoso", level=2)
    body(doc, "Se elimina un producto existente. El sistema confirma con [OK] Producto eliminado.")
    insert_image(doc, "10_eliminar_ok", "Eliminar producto exitoso")

    heading(doc, "5.9 Eliminar producto — caso de error (ID no existe)", level=2)
    body(doc, "Se intenta eliminar un producto con ID inexistente. El sistema responde con [ERROR].")
    insert_image(doc, "11_eliminar_error", "Eliminar producto — ID inexistente")

    # ── 6. PRUEBAS UNITARIAS ─────────────────────────────────────────────── #
    doc.add_page_break()
    heading(doc, "6. Pruebas Unitarias")
    body(doc,
         "Se implementaron 22 pruebas unitarias con el módulo unittest de Python, "
         "agrupadas en cuatro clases (una por operación CRUD). Cada clase tiene al menos "
         "un caso exitoso y uno de error.")

    tbl4 = doc.add_table(rows=5, cols=4); tbl4.style = "Table Grid"
    for j, h in enumerate(["Operación", "Exitosas", "Error", "Total"]):
        set_cell_bg(tbl4.rows[0].cells[j], "D9E1F2")
        tbl4.rows[0].cells[j].paragraphs[0].add_run(h).bold = True
    resumen = [("Crear","2","3","5"),("Leer","3","2","5"),
               ("Actualizar","3","4","7"),("Eliminar","3","2","5")]
    for i, row in enumerate(resumen, 1):
        for j, val in enumerate(row):
            tbl4.rows[i].cells[j].paragraphs[0].add_run(val)

    doc.add_paragraph()
    heading(doc, "6.1 Código fuente — test_productos.py", level=2)
    add_code_block(doc, TESTS_PY)

    heading(doc, "6.2 Resultado de ejecución de los tests", level=2)
    body(doc, "Resultado al ejecutar: python -m unittest test_productos -v")
    insert_image(doc, "12_tests", "22 pruebas unitarias — Ran 22 tests in 0.00s OK")

    # ── 7. CALIDAD DE CÓDIGO — SONARQUBE ─────────────────────────────────── #
    doc.add_page_break()
    heading(doc, "7. Prueba de Calidad del Código — SonarCloud")
    body(doc,
         "Se realizó el análisis de calidad del código utilizando SonarCloud "
         "(versión online de SonarQube), integrado con el repositorio de GitHub. "
         "El análisis evalúa fiabilidad, seguridad, mantenibilidad, cobertura y duplicaciones.")
    insert_image(doc, "13_sonarqube", "Dashboard de resultados en SonarCloud")

    # ── 8. GITHUB ────────────────────────────────────────────────────────── #
    heading(doc, "8. Control de Versiones — GitHub")
    body(doc,
         "El código fuente fue publicado en un repositorio de GitHub y compartido con "
         "el docente al correo jonathansanchez2948@correo.itm.edu.co.")

    tbl5 = doc.add_table(rows=2, cols=2); tbl5.style = "Table Grid"
    set_cell_bg(tbl5.rows[0].cells[0], "D9E1F2")
    set_cell_bg(tbl5.rows[0].cells[1], "D9E1F2")
    tbl5.rows[0].cells[0].paragraphs[0].add_run("Campo").bold = True
    tbl5.rows[0].cells[1].paragraphs[0].add_run("Valor").bold = True
    tbl5.rows[1].cells[0].paragraphs[0].add_run("URL del repositorio")
    tbl5.rows[1].cells[1].paragraphs[0].add_run("[PEGAR URL DE GITHUB AQUÍ]")

    doc.add_paragraph()
    insert_image(doc, "14_github", "Repositorio GitHub — archivos del proyecto")

    # ── 9. BIBLIOGRAFÍA ──────────────────────────────────────────────────── #
    doc.add_page_break()
    heading(doc, "9. Bibliografía")
    refs = [
        "Python Software Foundation. (2024). Python 3 Documentation. https://docs.python.org/3/",
        "Python Software Foundation. (2024). unittest — Unit testing framework. "
        "https://docs.python.org/3/library/unittest.html",
        "Python Software Foundation. (2024). json — JSON encoder and decoder. "
        "https://docs.python.org/3/library/json.html",
        "Martin, R. C. (2009). Clean Code: A Handbook of Agile Software Craftsmanship. Prentice Hall.",
        "Lutz, M. (2013). Learning Python (5th ed.). O'Reilly Media.",
        "SonarSource. (2024). SonarCloud Documentation. https://docs.sonarcloud.io/",
        "GitHub, Inc. (2024). GitHub Docs. https://docs.github.com/",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Bullet")

    out = "Guia4_Entrega_CRUD_Python.docx"
    doc.save(out)

    # reporte de imágenes encontradas / faltantes
    slots = [
        ("01_estructura","Estructura proyecto"), ("02_menu","Menú principal"),
        ("03_crear_ok","Crear OK"),              ("04_crear_error","Crear error"),
        ("05_leer_ok","Leer OK"),
        ("07_listar","Listar todos"),            ("08_actualizar_ok","Actualizar OK"),
        ("09_actualizar_error","Actualizar error"),("10_eliminar_ok","Eliminar OK"),
        ("11_eliminar_error","Eliminar error"),  ("12_tests","22 tests OK"),
        ("13_sonarqube","SonarCloud"),           ("14_github","GitHub"),
    ]
    encontradas = [s for s,_ in slots if find_image(s)]
    faltantes   = [f"{s}.png  →  {d}" for s,d in slots if not find_image(s)]

    print(f"\n[OK] Documento generado: {out}")
    print(f"     Imágenes insertadas : {len(encontradas)}/{len(slots)}")
    if faltantes:
        print("\n[!] Capturas faltantes (guardar en Capturas/):")
        for f in faltantes:
            print(f"    {f}")
    else:
        print("     Todas las capturas fueron insertadas correctamente.")


if __name__ == "__main__":
    build_document()
